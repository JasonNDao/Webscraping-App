from bs4 import BeautifulSoup as scrape
from docx import *
import requests
from PIL import Image
from io import BytesIO
import os



doc=Document()
stuff=requests.get("https://appsliced.co/apps/iphone?sort=latest&threshold=all&price=free&l=nav")
soup=scrape(stuff.text,"lxml")
l=soup.find("article",class_="app app_list first_app touch")
o=l.div.div.div.div.a
s=o.img
requestemp=requests.get(s.attrs["src"])
temp=Image.open(BytesIO(requestemp.content))
try:                                       #if no error
    temp.save("./App"+"."+temp.format,temp.format)          #save image(name, format) name must be like ray.jpg
except IOError:                            #error what happens
    a=6
doc.add_picture("App."+str(temp.format))
j=doc.add_paragraph("")
k=j.add_run(o.attrs["title"])
k.bold=True
doc.add_paragraph(o.attrs["href"])
o=o.parent
o=o.find_next_sibling()
o=o.a.div
o=o.find_next_sibling().div.text
doc.add_paragraph(o)

for article in soup.findAll("article",class_="app app_list touch"):
    o=article.div.div.div.div.a
    s=o.img
    requestemp=requests.get(s.attrs["src"])
    temp=Image.open(BytesIO(requestemp.content))
    try:                                       #if no error
        temp.save("./App"+"."+temp.format,temp.format)          #save image(name, format) name must be like ray.jpg
    except IOError:                            #error what happens
        a=6
    doc.add_paragraph("")
    doc.add_picture("App."+str(temp.format))
    j=doc.add_paragraph("")
    k=j.add_run(o.attrs["title"])
    k.bold=True
    doc.add_paragraph(o.attrs["href"])
    o=o.parent
    o=o.find_next_sibling()
    o=o.a.div
    o=o.find_next_sibling().div.text
    doc.add_paragraph(o)
doc.save("FreeApps.docx")
os.startfile("FreeApps.docx")

    


